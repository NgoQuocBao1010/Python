import tkinter as tk
import docx

doc = docx.Document()
doc.save('lastalgorithm.docx')

lst_of_ele = []


def do_operator(x1, symbol=None, x2=0):
	if symbol == '+':
		return x1 + x2
	elif symbol == 'x':
		return x1 * x2
	elif symbol == '-':
		return x1 - x2
	elif symbol == '/':
		return x1 / x2
	elif symbol == None:
		return x1
	else:
		return 'Invalid Error'


def check_for_operand():
	global lst_of_ele
	time = ['x', '/']
	for index in range(len(lst_of_ele)):
		if lst_of_ele[index] in time:
			return index
	return -1


def do_the_math():
	global lst_of_ele
	
	while check_for_operand() != -1:
		index = check_for_operand()
		x1 = lst_of_ele.pop(index-1)
		sym = lst_of_ele.pop(index-1)
		x2 = lst_of_ele.pop(index-1)
		result = do_operator(x1, sym, x2)
		lst_of_ele.insert(index-1, result)

	while len(lst_of_ele) > 1:
		x1 = lst_of_ele.pop(0)
		sym = lst_of_ele.pop(0)
		x2 = lst_of_ele.pop(0)
		result = do_operator(x1, sym, x2)
		lst_of_ele.insert(0, result)

	return lst_of_ele[0]
	

def convert_to_list(algorithm):
	global lst_of_ele
	number = '0123456789'
	operand = '+-x/'
	var = ''
	lst_of_ele = []
	for element in algorithm:
		if element not in number:
			lst_of_ele.append(var)
			lst_of_ele.append(element)
			var = ''
		else:
			var += element
	lst_of_ele.append(var)

	for index in range(len(lst_of_ele)):
		if lst_of_ele[index] not in operand or lst_of_ele[index] == '':
			try:
				lst_of_ele[index] = int(lst_of_ele[index])
			except:
				return '***ERROR 404***'
	print(lst_of_ele)
	return do_the_math()


def summit():
	result = 0
	algorithm = main_al.get()
	clear_documet()
	docu = docx.Document('lastalgorithm.docx')
	docu.add_paragraph(algorithm)
	docu.save('lastalgorithm.docx')
	result = convert_to_list(algorithm)

	main_al.delete(0, tk.END)
	main_al.insert(tk.END,result)
	main_al['justify'] = 'right'


def clear_main():
	if main_al['justify'] == 'right':
		main_al['justify'] = 'left'
	main_al.delete(0, tk.END)


def clear_documet():
	clear_main()
	doc = docx.Document()
	doc.save('lastalgorithm.docx')


def last_algorithm():
	doc = docx.Document('lastalgorithm.docx')
	text = doc.paragraphs[0].text
	clear_main()
	main_al.insert(0, text)

class button():
	def __init__(self,name, text, row, column, x, y):
		self.name = name
		self.text = text
		self.row = row
		self.column = column
		self.x = x
		self.y = y
		self.create()

	def create(self):
		self.name = tk.Button(num_lab, text=self.text, bg='black', fg='white', 
						font=("Helvetica", 14, 'bold'), command=self.insert)
		self.name.grid(row=self.row, column=self.column, ipadx=self.x, ipady=self.y)

	def insert(self):
		if main_al['justify'] == 'right':
			main_al['justify'] = 'left'
			main_al.delete(0, tk.END)
		main_al.insert(tk.END, self.name['text'])


root = tk.Tk()
root.title('Calculator')

canvas = tk.Canvas(root, height=500, width=500, bg='gray')
canvas.pack()

main_al = tk.Entry(canvas, bg='white', font=("Courier", 20), selectbackground='red',
					highlightcolor='yellow', bd=5)
main_al.place(relx=0.05, rely=0.05, relwidth=0.9, relheight=0.1)

summit = tk.Button(canvas, text='=', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command=summit)
summit.place(relx=0.75, rely=0.2, relwidth=0.2, relheight=0.1)

delete = tk.Button(canvas, text='Del', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command=lambda: main_al.delete(len(main_al.get()) - 1))
delete.place(relx=0.65, rely=0.3, relwidth=0.15, relheight=0.1)

clear = tk.Button(canvas, text='C', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command=clear_main)
clear.place(relx=0.8, rely=0.3, relwidth=0.15, relheight=0.1)

reset = tk.Button(canvas, text='Reset', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command=clear_documet)
reset.place(relx=0.8, rely=0.4, relwidth=0.15, relheight=0.1)

L_A = tk.Button(canvas, text='L.A', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command=last_algorithm)
L_A.place(relx=0.65, rely=0.4, relwidth=0.15, relheight=0.1)

plus = tk.Button(canvas, text='+', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command= lambda: main_al.insert(tk.END, plus['text']))
plus.place(relx=0.8, rely=0.55, relwidth=0.15, relheight=0.1)

minus = tk.Button(canvas, text='-', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command= lambda: main_al.insert(tk.END, minus['text']))
minus.place(relx=0.65, rely=0.55, relwidth=0.15, relheight=0.1)

time = tk.Button(canvas, text='x', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command= lambda: main_al.insert(tk.END, time['text']))
time.place(relx=0.8, rely=0.65, relwidth=0.15, relheight=0.1)

divide = tk.Button(canvas, text='/', bg='black', fg='white', font=("Gothic", 15, 'bold'),
					command= lambda: main_al.insert(tk.END, divide['text']))
divide.place(relx=0.65, rely=0.65, relwidth=0.15, relheight=0.1)

num_lab = tk.Frame(canvas, bg='black')
num_lab.place(relx=0.05, rely=0.2, relwidth=0.58, relheight=0.75)

text = 1
for row in range(1, 4):
	for column in range(1, 4):
		bt = button('button', text, row, column, 35, 30)
		text += 1

bt = button('button', 0, 4, 2, 35, 30)


root.mainloop()

# This is Unfinished